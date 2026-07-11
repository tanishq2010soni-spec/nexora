import io
import uuid
from typing import List, Optional
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.application.interfaces.embedding_service import EmbeddingService, VectorDBRepository
from src.infrastructure.database.models import Document, KnowledgeBase
from src.infrastructure.logging.logger import get_logger

logger = get_logger(__name__)


class DocumentService:
    def __init__(self, embedding_service: EmbeddingService, vector_db: VectorDBRepository):
        self.embedding_service = embedding_service
        self.vector_db = vector_db

    async def ingest_document(
        self,
        db: AsyncSession,
        kb_id: uuid.UUID,
        filename: str,
        file_bytes: bytes,
        org_id: uuid.UUID,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> Document:
        """
        Parses text, chunks it, generates vector embeddings, writes to Qdrant,
        and logs metadata records in PostgreSQL.
        """
        logger.info("Ingesting document text", filename=filename, kb_id=str(kb_id))
        
        # 1. Text extraction based on file extension
        text = self._extract_text(filename, file_bytes)
        if not text.strip():
            raise ValueError(f"Extracted document text in {filename} is empty.")

        # 2. Text splitting / chunking
        chunks = self._chunk_text(text, chunk_size, chunk_overlap)
        
        # 3. Create document record in database (pending state)
        db_doc = Document(
            id=uuid.uuid4(),
            kb_id=kb_id,
            filename=filename,
            file_type=filename.split(".")[-1].lower(),
            file_size=len(file_bytes),
            storage_path=f"kb/{kb_id}/{filename}",
            status="processing",
            chunk_count=len(chunks),
        )
        db.add(db_doc)
        await db.commit()
        await db.refresh(db_doc)

        try:
            # 4. Generate embeddings and upload to Qdrant vector database
            points = []
            for idx, chunk in enumerate(chunks):
                vector = await self.embedding_service.generate_embedding(chunk)
                points.append({
                    "id": str(uuid.uuid4()),
                    "vector": vector,
                    "text": chunk,
                    "document_id": str(db_doc.id),
                    "metadata": {
                        "chunk_index": idx,
                        "kb_id": str(kb_id),
                        "filename": filename
                    }
                })

            # Upsert into Qdrant vector index
            await self.vector_db.upsert_chunks(str(org_id), points)
            
            # Update status to active
            db_doc.status = "active"
            await db.commit()
            logger.info("Document ingestion completed successfully", doc_id=str(db_doc.id))
            return db_doc
        except Exception as e:
            logger.error("Failed to complete document ingestion. Marking database status as failed.", error=str(e))
            db_doc.status = "failed"
            await db.commit()
            raise

    async def delete_document(self, db: AsyncSession, doc_id: uuid.UUID, org_id: uuid.UUID) -> bool:
        """
        Removes vector references from Qdrant and database records.
        """
        stmt = select(Document).where(Document.id == doc_id)
        result = await db.execute(stmt)
        doc = result.scalar_one_or_none()
        if not doc:
            return False

        # Delete from Qdrant vector index
        await self.vector_db.delete_by_document(str(org_id), str(doc_id))

        # Delete database model record
        await db.delete(doc)
        await db.commit()
        logger.info("Deleted document and index points", doc_id=str(doc_id))
        return True

    def _extract_text(self, filename: str, file_bytes: bytes) -> str:
        ext = filename.split(".")[-1].lower()
        if ext == "txt":
            return file_bytes.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            pdf = PdfReader(io.BytesIO(file_bytes))
            text_runs = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_runs.append(t)
            return "\n".join(text_runs)
        elif ext in ["docx", "doc"]:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_runs = []
            for para in doc.paragraphs:
                if para.text:
                    text_runs.append(para.text)
            return "\n".join(text_runs)
        else:
            raise ValueError(f"Unsupported document file extension: {ext}")

    def _chunk_text(self, text: str, size: int, overlap: int) -> List[str]:
        if not text:
            return []
        
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + size, text_len)
            chunk = text[start:end]
            chunks.append(chunk)
            
            # Step forward by (size - overlap)
            start += size - overlap
            if start >= text_len or size <= overlap:
                break
                
        return chunks
